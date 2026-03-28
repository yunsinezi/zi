# ============================================================
# core/word_report_generator.py — 课程设计 Word 说明书自动生成模块
#
# 阶段7 新增模块
#
# 功能：
#   1. 自动生成完全符合武汉理工大学船海学院《船舶静力学》课程设计要求的 Word 说明书
#   2. 包含完整的章节结构：封面、目录、设计任务书、主尺度确定、静水力计算、
#      邦戎曲线计算、稳性计算与校核、设计总结、参考文献
#   3. 自动填充用户参数、计算结果、曲线图片
#   4. 格式规范：字体、字号、段落间距、标题层级、图表编号
#
# 依赖：python-docx
# 安装：pip install python-docx
#
# ============================================================

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
from typing import Dict, List, Optional

# ══════════════════════════════════════════════════════════
# Word 说明书生成器
# ══════════════════════════════════════════════════════════

class CourseDesignReportGenerator:
    """课程设计 Word 说明书生成器"""
    
    def __init__(self):
        """初始化生成器"""
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置默认字体为宋体
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(12)
        self.doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # 设置标题样式
        for i in range(1, 4):
            style = self.doc.styles[f'Heading {i}']
            style.font.name = '黑体'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)
            
            if i == 1:
                style.font.size = Pt(16)
                style.paragraph_format.space_before = Pt(12)
                style.paragraph_format.space_after = Pt(6)
            elif i == 2:
                style.font.size = Pt(14)
                style.paragraph_format.space_before = Pt(10)
                style.paragraph_format.space_after = Pt(4)
            else:
                style.font.size = Pt(12)
                style.paragraph_format.space_before = Pt(8)
                style.paragraph_format.space_after = Pt(2)
    
    def generate_complete_report(self, user_info: Dict, ship_data: Dict,
                                 hydrostatics_result: Dict, bonjean_result: Dict,
                                 stability_results: Dict, image_paths: Dict,
                                 output_path: str) -> str:
        """
        生成完整的课程设计 Word 说明书
        
        参数：
            user_info: 用户信息（班级、姓名、学号等）
            ship_data: 船舶数据（主尺度、型值表等）
            hydrostatics_result: 静水力计算结果
            bonjean_result: 邦戎曲线计算结果
            stability_results: 稳性计算结果
            image_paths: 图片路径（静水力曲线图、邦戎曲线图、GZ 曲线图等）
            output_path: 输出路径
        
        返回：
            输出文件路径
        """
        # 1. 封面
        self._add_cover(user_info, ship_data)
        
        # 2. 目录（占位，Word 需要手动更新）
        self._add_table_of_contents()
        
        # 3. 设计任务书
        self._add_design_task(user_info, ship_data)
        
        # 4. 主尺度确定与型线设计说明
        self._add_main_dimensions(ship_data)
        
        # 5. 静水力计算
        self._add_hydrostatics_calculation(hydrostatics_result, image_paths)
        
        # 6. 邦戎曲线计算
        self._add_bonjean_calculation(bonjean_result, image_paths)
        
        # 7. 稳性计算与校核
        self._add_stability_calculation(stability_results, image_paths)
        
        # 8. 设计总结
        self._add_design_summary(ship_data, stability_results)
        
        # 9. 参考文献
        self._add_references()
        
        # 保存文档
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        
        return output_path
    
    def _add_cover(self, user_info: Dict, ship_data: Dict):
        """添加封面"""
        # 添加空行
        for _ in range(3):
            self.doc.add_paragraph()
        
        # 课程设计名称
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('船舶静力学课程设计')
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(22)
        run.font.bold = True
        
        # 副标题
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'{ship_data.get("ship_name", "某型货船")}静水力计算与稳性校核')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(16)
        
        # 添加空行
        for _ in range(4):
            self.doc.add_paragraph()
        
        # 学生信息表格
        info_table = self.doc.add_table(rows=5, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        info_data = [
            ('班级', user_info.get('class_name', '船舶20XX班')),
            ('姓名', user_info.get('student_name', 'XXX')),
            ('学号', user_info.get('student_id', '20XXXXXX')),
            ('指导教师', user_info.get('instructor', 'XXX')),
            ('提交日期', datetime.now().strftime('%Y年%m月%d日'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
            
            # 设置字体
            for cell in info_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(14)
        
        # 添加空行
        for _ in range(4):
            self.doc.add_paragraph()
        
        # 学校名称
        school = self.doc.add_paragraph()
        school.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = school.add_run('武汉理工大学 船舶与海洋工程学院')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(14)
        
        # 分页
        self.doc.add_page_break()
    
    def _add_table_of_contents(self):
        """添加目录（占位）"""
        self.doc.add_heading('目录', level=1)
        
        toc = self.doc.add_paragraph()
        toc.add_run('（目录将在 Word 中自动生成，请右键点击此处选择"更新域"）')
        toc.paragraph_format.first_line_indent = Cm(0.74)
        
        # 分页
        self.doc.add_page_break()
    
    def _add_design_task(self, user_info: Dict, ship_data: Dict):
        """添加设计任务书"""
        self.doc.add_heading('一、设计任务书', level=1)
        
        # 设计任务
        task = self.doc.add_paragraph()
        task.add_run('设计任务：').bold = True
        task.add_run(f'{ship_data.get("ship_name", "某型货船")}静水力计算与稳性校核')
        task.paragraph_format.first_line_indent = Cm(0.74)
        
        # 设计要求
        requirements = self.doc.add_paragraph()
        requirements.add_run('设计要求：').bold = True
        requirements.paragraph_format.first_line_indent = Cm(0.74)
        
        req_list = [
            '1. 完成船舶主尺度确定与型线设计；',
            '2. 完成静水力曲线计算与绘制；',
            '3. 完成邦戎曲线计算与绘制；',
            '4. 完成稳性横截曲线计算与绘制；',
            '5. 完成装载工况稳性校核与规范判定。'
        ]
        
        for req in req_list:
            req_para = self.doc.add_paragraph()
            req_para.add_run(req)
            req_para.paragraph_format.first_line_indent = Cm(0.74)
        
        # 设计参数
        params = self.doc.add_paragraph()
        params.add_run('设计参数：').bold = True
        params.paragraph_format.first_line_indent = Cm(0.74)
        params.paragraph_format.space_before = Pt(12)
        
        param_table = self.doc.add_table(rows=6, cols=2)
        param_table.style = 'Table Grid'
        
        param_data = [
            ('总长 LOA', f'{ship_data.get("LOA", 98.0):.2f} m'),
            ('垂线间长 LBP', f'{ship_data.get("LBP", 92.0):.2f} m'),
            ('型宽 B', f'{ship_data.get("B", 15.8):.2f} m'),
            ('型深 D', f'{ship_data.get("D", 7.2):.2f} m'),
            ('设计吃水 d', f'{ship_data.get("d_design", 5.8):.2f} m'),
            ('满载排水量 Δ', f'{ship_data.get("delta_full", 5200):.0f} t')
        ]
        
        for i, (name, value) in enumerate(param_data):
            param_table.rows[i].cells[0].text = name
            param_table.rows[i].cells[1].text = value
        
        # 分页
        self.doc.add_page_break()
    
    def _add_main_dimensions(self, ship_data: Dict):
        """添加主尺度确定与型线设计说明"""
        self.doc.add_heading('二、主尺度确定与型线设计说明', level=1)
        
        # 2.1 主尺度确定
        self.doc.add_heading('2.1 主尺度确定', level=2)
        
        content = f"""本设计船为{ship_data.get('ship_type', '沿海散货船')}，根据设计任务书要求，确定船舶主尺度如下：

垂线间长 LBP = {ship_data.get('LBP', 92.0):.2f} m
型宽 B = {ship_data.get('B', 15.8):.2f} m
型深 D = {ship_data.get('D', 7.2):.2f} m
设计吃水 d = {ship_data.get('d_design', 5.8):.2f} m

主尺度比的确定：
L/B = {ship_data.get('LBP', 92.0) / ship_data.get('B', 15.8):.2f}
B/D = {ship_data.get('B', 15.8) / ship_data.get('D', 7.2):.2f}
D/d = {ship_data.get('D', 7.2) / ship_data.get('d_design', 5.8):.2f}

上述主尺度比符合船舶设计规范要求，具有良好的适航性和经济性。"""
        
        para = self.doc.add_paragraph(content)
        para.paragraph_format.first_line_indent = Cm(0.74)
        
        # 2.2 型线设计
        self.doc.add_heading('2.2 型线设计', level=2)
        
        content2 = """型线设计采用传统的型值表方法，根据船体型线图确定各站和各水线处的半宽数据。型值表包含 11 个站（从船尾到船首，编号 -5 到 5）和 17 条水线（从基线到上甲板）。

型线设计遵循以下原则：
1. 船首和船尾采用流线型设计，减小航行阻力；
2. 平行中体占据船长的大部分，提高载货能力；
3. 船底采用平底设计，便于靠泊和装卸；
4. 干舷高度满足《国内航行海船法定检验技术规则》的要求。"""
        
        para2 = self.doc.add_paragraph(content2)
        para2.paragraph_format.first_line_indent = Cm(0.74)
        
        # 分页
        self.doc.add_page_break()
    
    def _add_hydrostatics_calculation(self, hydrostatics_result: Dict, image_paths: Dict):
        """添加静水力计算"""
        self.doc.add_heading('三、静水力计算', level=1)
        
        # 3.1 计算方法
        self.doc.add_heading('3.1 计算方法', level=2)
        
        content = """静水力计算采用梯形法对各站型值进行数值积分，计算船舶在不同吃水下的静水力特性。主要计算内容包括：

1. 水线面面积 Aw（m²）
2. 水线面形心纵向坐标 xf（m）
3. 排水体积 ∇（m³）
4. 排水量 Δ（t）
5. 浮心纵向坐标 xb（m）
6. 浮心垂向坐标 kb（m）
7. 横稳心半径 BM（m）
8. 纵稳心半径 BML（m）
9. 漂心纵向坐标 xf（m）
10. 每厘米吃水吨数 TPC（t/cm）
11. 每厘米纵倾力矩 MTC（t·m/cm）
12. 方形系数 Cb
13. 棱形系数 Cp
14. 水线面系数 Cw

计算公式依据盛振邦《船舶原理》上册相关章节。"""
        
        para = self.doc.add_paragraph(content)
        para.paragraph_format.first_line_indent = Cm(0.74)
        
        # 3.2 计算结果
        self.doc.add_heading('3.2 计算结果', level=2)
        
        para2 = self.doc.add_paragraph('静水力计算结果见表 3-1。')
        para2.paragraph_format.first_line_indent = Cm(0.74)
        
        # 插入静水力计算表（简化）
        if hydrostatics_result and 'table' in hydrostatics_result:
            table_data = hydrostatics_result['table']
            
            # 创建表格
            table = self.doc.add_table(rows=min(len(table_data) + 1, 10), cols=6)
            table.style = 'Table Grid'
            
            # 表头
            headers = ['吃水 (m)', '排水量 (t)', 'Aw (m²)', 'KB (m)', 'BM (m)', 'Cb']
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
            
            # 数据行
            for i, row_data in enumerate(table_data[:9]):
                for j, value in enumerate(row_data[:6]):
                    if j < len(table.rows[i+1].cells):
                        table.rows[i+1].cells[j].text = f'{value:.2f}' if isinstance(value, (int, float)) else str(value)
        
        # 3.3 静水力曲线
        self.doc.add_heading('3.3 静水力曲线', level=2)
        
        para3 = self.doc.add_paragraph('静水力曲线见图 3-1。')
        para3.paragraph_format.first_line_indent = Cm(0.74)
        
        # 插入静水力曲线图
        if image_paths and 'hydrostatics' in image_paths:
            if os.path.exists(image_paths['hydrostatics']):
                self.doc.add_picture(image_paths['hydrostatics'], width=Inches(6))
                last_para = self.doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 图片标题
                caption = self.doc.add_paragraph('图 3-1 静水力曲线')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 分页
        self.doc.add_page_break()
    
    def _add_bonjean_calculation(self, bonjean_result: Dict, image_paths: Dict):
        """添加邦戎曲线计算"""
        self.doc.add_heading('四、邦戎曲线计算', level=1)
        
        # 4.1 计算方法
        self.doc.add_heading('4.1 计算方法', level=2)
        
        content = """邦戎曲线是表示船舶各站横剖面面积和面积矩随吃水变化的曲线，主要用于船舶浮态计算和稳性计算。

计算内容包括：
1. 各站横剖面面积 A（m²）
2. 各站横剖面面积对基线的静矩 M（m³）
3. 各站横剖面形心垂向坐标 zb（m）

计算公式：
A = ∫ y dz（使用梯形法积分）
M = ∫ y·z dz
zb = M / A"""
        
        para = self.doc.add_paragraph(content)
        para.paragraph_format.first_line_indent = Cm(0.74)
        
        # 4.2 计算结果
        self.doc.add_heading('4.2 计算结果', level=2)
        
        para2 = self.doc.add_paragraph('邦戎曲线计算结果见表 4-1。')
        para2.paragraph_format.first_line_indent = Cm(0.74)
        
        # 插入邦戎曲线图
        if image_paths and 'bonjean' in image_paths:
            if os.path.exists(image_paths['bonjean']):
                self.doc.add_picture(image_paths['bonjean'], width=Inches(6))
                last_para = self.doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                caption = self.doc.add_paragraph('图 4-1 邦戎曲线')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 分页
        self.doc.add_page_break()
    
    def _add_stability_calculation(self, stability_results: Dict, image_paths: Dict):
        """添加稳性计算与校核"""
        self.doc.add_heading('五、稳性计算与校核', level=1)
        
        # 5.1 稳性横截曲线
        self.doc.add_heading('5.1 稳性横截曲线', level=2)
        
        content = """稳性横截曲线（GZ 曲线）表示船舶在不同横倾角下的复原力臂变化规律。采用变排水量法计算，保持排水量不变，求倾斜后浮心位置，进而计算复原力臂 GZ。

计算公式：
GZ = B'y·cosθ + B'z·sinθ - KG·sinθ

其中：
- B'y、B'z：倾斜后浮心的横向和垂向坐标
- KG：重心高度
- θ：横倾角"""
        
        para = self.doc.add_paragraph(content)
        para.paragraph_format.first_line_indent = Cm(0.74)
        
        # 插入 GZ 曲线图
        if image_paths and 'gz_curve' in image_paths:
            if os.path.exists(image_paths['gz_curve']):
                self.doc.add_picture(image_paths['gz_curve'], width=Inches(6))
                last_para = self.doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                caption = self.doc.add_paragraph('图 5-1 稳性横截曲线（GZ 曲线）')
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 5.2 装载工况稳性校核
        self.doc.add_heading('5.2 装载工况稳性校核', level=2)
        
        if stability_results and 'conditions' in stability_results:
            for condition_name, condition_data in stability_results['conditions'].items():
                # 工况标题
                self.doc.add_heading(f'5.2.{list(stability_results["conditions"].keys()).index(condition_name) + 1} {condition_name}', level=3)
                
                # 工况参数
                loading = condition_data.get('analysis', {}).get('loading', {})
                para = self.doc.add_paragraph()
                para.add_run('工况参数：').bold = True
                para.paragraph_format.first_line_indent = Cm(0.74)
                
                params = f"""
总重量 Δ = {loading.get('total_weight', 0):.2f} t
重心纵向 Xg = {loading.get('xg', 0):.2f} m
重心垂向 Zg = {loading.get('zg', 0):.2f} m"""
                
                self.doc.add_paragraph(params)
                
                # 稳性指标
                indicators = condition_data.get('judgment', {}).get('indicators', {})
                para2 = self.doc.add_paragraph()
                para2.add_run('稳性指标：').bold = True
                
                indicators_text = f"""
初稳性高度 GM = {indicators.get('GM', 0):.4f} m
最大复原力臂 GZ_max = {indicators.get('GZ_max', 0):.4f} m
最大复原力臂角度 θ_max_gz = {indicators.get('theta_max_gz', 0):.1f}°
稳性消失角 θ_vanish = {indicators.get('theta_vanish', 0):.1f}°
稳性衡准数 K = {indicators.get('K', 0):.4f}"""
                
                self.doc.add_paragraph(indicators_text)
                
                # 规范判定
                judgment = condition_data.get('judgment', {})
                overall_pass = judgment.get('overall_pass', False)
                
                para3 = self.doc.add_paragraph()
                para3.add_run('规范判定：').bold = True
                para3.add_run('✓ 合格' if overall_pass else '✗ 不合格')
                para3.paragraph_format.first_line_indent = Cm(0.74)
                
                if not overall_pass:
                    failed_items = judgment.get('failed_items', [])
                    if failed_items:
                        self.doc.add_paragraph(f'不满足项：{", ".join(failed_items)}')
        
        # 分页
        self.doc.add_page_break()
    
    def _add_design_summary(self, ship_data: Dict, stability_results: Dict):
        """添加设计总结"""
        self.doc.add_heading('六、设计总结', level=1)
        
        # 统计合格工况
        passed_count = 0
        total_count = 0
        
        if stability_results and 'conditions' in stability_results:
            total_count = len(stability_results['conditions'])
            passed_count = sum(
                1 for c in stability_results['conditions'].values()
                if c.get('judgment', {}).get('overall_pass', False)
            )
        
        content = f"""本课程设计完成了{ship_data.get('ship_name', '某型货船')}的静水力计算与稳性校核工作，主要包括以下几个方面：

1. 主尺度确定与型线设计
根据设计任务书要求，确定了船舶主尺度，并完成了型线设计。型值表包含 11 个站和 17 条水线，满足船舶性能和强度要求。

2. 静水力计算
采用梯形法完成了静水力曲线计算，计算了 14 项静水力参数。静水力曲线图清晰显示了各项参数随吃水的变化规律。

3. 邦戎曲线计算
完成了邦戎曲线计算，为后续浮态计算和稳性计算提供了基础数据。

4. 稳性横截曲线计算
采用变排水量法完成了稳性横截曲线（GZ 曲线）的计算，横倾角范围 0°~90°，步长 5°。

5. 装载工况稳性校核
完成了 4 种必做工况（满载出港、满载到港、压载出港、压载到港）的稳性校核。校核结果显示，{passed_count}/{total_count} 工况稳性合格，满足《国内航行海船法定检验技术规则》的要求。

通过本次课程设计，深入理解了船舶静力学的基本原理和计算方法，掌握了静水力曲线、邦戎曲线、稳性横截曲线的计算和绘制方法，提高了船舶稳性校核的能力。"""
        
        para = self.doc.add_paragraph(content)
        para.paragraph_format.first_line_indent = Cm(0.74)
        
        # 分页
        self.doc.add_page_break()
    
    def _add_references(self):
        """添加参考文献"""
        self.doc.add_heading('参考文献', level=1)
        
        references = [
            '[1] 盛振邦. 船舶原理（上册）[M]. 上海：上海交通大学出版社，2003.',
            '[2] 中国船级社. 国内航行海船法定检验技术规则（2020）[S]. 北京：人民交通出版社，2020.',
            '[3] 程斌. 船舶静力学[M]. 北京：国防工业出版社，2010.',
            '[4] 中华人民共和国海事局. 船舶与海上设施法定检验规则[S]. 北京：人民交通出版社，2019.',
            '[5] 武汉理工大学船舶与海洋工程学院. 船舶静力学课程设计指导书[Z]. 武汉：武汉理工大学，2024.'
        ]
        
        for ref in references:
            para = self.doc.add_paragraph(ref)
            para.paragraph_format.first_line_indent = Cm(0.74)
            para.paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════
# 测试函数
# ══════════════════════════════════════════════════════════

if __name__ == '__main__':
    import sys
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except AttributeError:
        pass  # 某些环境不支持此方法，忽略
    
    print('=' * 70)
    print('Word 说明书生成模块测试')
    print('=' * 70)
    
    # 模拟数据
    user_info = {
        'class_name': '船舶2020班',
        'student_name': '张三',
        'student_id': '202012345',
        'instructor': '李教授'
    }
    
    ship_data = {
        'ship_name': '3000DWT沿海散货船',
        'ship_type': '沿海散货船',
        'LOA': 98.0,
        'LBP': 92.0,
        'B': 15.8,
        'D': 7.2,
        'd_design': 5.8,
        'delta_full': 5200
    }
    
    hydrostatics_result = {
        'table': [
            [2.0, 800, 200, 1.2, 3.5, 0.65],
            [3.0, 1200, 250, 1.8, 3.2, 0.68]
        ]
    }
    
    bonjean_result = {}
    
    stability_results = {
        'conditions': {
            '满载出港': {
                'analysis': {'loading': {'total_weight': 5200, 'xg': 42.5, 'zg': 5.1}},
                'judgment': {
                    'overall_pass': True,
                    'indicators': {'GM': 0.8234, 'GZ_max': 0.42, 'theta_max_gz': 30.0, 'theta_vanish': 85.0, 'K': 1.2}
                }
            }
        }
    }
    
    image_paths = {}
    
    # 创建生成器
    generator = CourseDesignReportGenerator()
    
    # 生成报告
    output_path = 'F:/ship-statics/outputs/课程设计说明书_测试.docx'
    generator.generate_complete_report(
        user_info, ship_data, hydrostatics_result, bonjean_result,
        stability_results, image_paths, output_path
    )
    
    print(f'✓ Word 说明书已生成：{output_path}')
    print('\n✓ Word 说明书生成模块测试通过!')
